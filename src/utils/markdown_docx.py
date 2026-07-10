from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

HASH_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
BOLD_LINE = re.compile(r"^\*\*(.+?)\*\*$")
ARTICLE_PATTERN = re.compile(
    r"^(ĐIỀU|DIEU|ĐIÊU|MỤC|MUC|CHƯƠNG|CHUONG|PHẦN|PHAN|"
    r"ARTICLE|SECTION|CHAPTER|CLAUSE)\b",
    re.IGNORECASE,
)
ARTICLE_NUMBERED = re.compile(
    r"^(Điều|Dieu|Mục|Muc|Chương|Chuong|Article|Section|Clause)\s+\d+",
    re.IGNORECASE,
)
BULLET_PATTERN = re.compile(r"^[-*•]\s+(.*)$")
NUMBERED_PATTERN = re.compile(
    r"^(\d+[.)]|[a-zA-Z][.)]|[ivxlcdm]+[.)])\s+(.*)$",
    re.IGNORECASE,
)
SIGNATURE_HINT = re.compile(
    r"(đại diện|dai dien|ký tên|ky ten|chữ ký|chu ky|"
    r"signature|signed by|party\s*[ab]|bên\s*[ab]|ben\s*[ab])",
    re.IGNORECASE,
)


def _set_run_font(
    run,
    *,
    bold: bool = False,
    size_pt: float = 12,
    italic: bool = False,
) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _strip_inline_markdown(text: str) -> str:
    cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    cleaned = re.sub(r"__(.+?)__", r"\1", cleaned)
    cleaned = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", cleaned)
    cleaned = re.sub(r"`([^`]+)`", r"\1", cleaned)
    return cleaned.replace("**", "").replace("__", "").strip()


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)


def _add_paragraph(
    doc: Document,
    text: str,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    bold: bool = False,
    size_pt: float = 12,
    space_before: float = 0,
    space_after: float = 6,
    first_line_indent_cm: float | None = None,
    keep_with_next: bool = False,
):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line_indent_cm is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    if keep_with_next:
        paragraph.paragraph_format.keep_with_next = True

    run = paragraph.add_run(text)
    _set_run_font(run, bold=bold, size_pt=size_pt)
    return paragraph


def _is_title_candidate(text: str, index: int) -> bool:
    if index > 2:
        return False
    upper = text.upper()
    return (
        "HỢP ĐỒNG" in upper
        or "HOP DONG" in upper
        or "CONTRACT" in upper
        or "AGREEMENT" in upper
        or len(text) <= 120
    )


def _is_article_heading(text: str) -> bool:
    return bool(ARTICLE_PATTERN.match(text) or ARTICLE_NUMBERED.match(text))


def _parse_lines(markdown_text: str) -> list[dict]:
    items: list[dict] = []

    for raw in markdown_text.splitlines():
        stripped = raw.strip()
        if not stripped:
            if items and items[-1]["kind"] != "blank":
                items.append({"kind": "blank", "text": "", "bold": False})
            continue

        hash_match = HASH_HEADING.match(stripped)
        if hash_match:
            level = len(hash_match.group(1))
            text = _strip_inline_markdown(hash_match.group(2))
            if text:
                items.append(
                    {
                        "kind": "title" if level == 1 else "heading",
                        "text": text,
                        "bold": True,
                    }
                )
            continue

        bold_match = BOLD_LINE.match(stripped)
        if bold_match:
            text = _strip_inline_markdown(bold_match.group(1))
            if text:
                kind = "heading" if _is_article_heading(text) or len(text) <= 120 else "body"
                items.append({"kind": kind, "text": text, "bold": True})
            continue

        bullet_match = BULLET_PATTERN.match(stripped)
        if bullet_match:
            text = _strip_inline_markdown(bullet_match.group(1))
            if text:
                items.append({"kind": "bullet", "text": text, "bold": False})
            continue

        numbered_match = NUMBERED_PATTERN.match(stripped)
        if numbered_match and not _is_article_heading(stripped):
            marker = numbered_match.group(1)
            body = _strip_inline_markdown(numbered_match.group(2))
            text = f"{marker} {body}".strip()
            if text:
                items.append({"kind": "numbered", "text": text, "bold": False})
            continue

        text = _strip_inline_markdown(stripped)
        if not text:
            continue

        if _is_article_heading(text):
            items.append({"kind": "heading", "text": text, "bold": True})
        elif SIGNATURE_HINT.search(text) and len(text) <= 80:
            items.append({"kind": "signature", "text": text, "bold": True})
        else:
            items.append({"kind": "body", "text": text, "bold": False})

    return items


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    """Convert contract text into a professionally formatted Word document."""
    doc = Document()
    _configure_document(doc)

    items = _parse_lines(markdown_text)
    title_written = False
    content_index = 0

    for item in items:
        kind = item["kind"]
        text = item["text"]

        if kind == "blank":
            continue

        if not title_written and (
            kind == "title" or (kind in {"heading", "body"} and _is_title_candidate(text, content_index))
        ):
            _add_paragraph(
                doc,
                text.upper(),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                bold=True,
                size_pt=14,
                space_before=0,
                space_after=18,
            )
            title_written = True
            content_index += 1
            continue

        if kind == "heading" or (item["bold"] and _is_article_heading(text)):
            _add_paragraph(
                doc,
                text,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                bold=True,
                size_pt=12,
                space_before=14,
                space_after=6,
                keep_with_next=True,
            )
            content_index += 1
            continue

        if kind == "bullet":
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(text)
            _set_run_font(run, size_pt=12)
            content_index += 1
            continue

        if kind == "numbered":
            _add_paragraph(
                doc,
                text,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent_cm=0.75,
                space_after=4,
            )
            content_index += 1
            continue

        if kind == "signature":
            _add_paragraph(
                doc,
                text,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                bold=True,
                size_pt=12,
                space_before=18,
                space_after=28,
            )
            content_index += 1
            continue

        _add_paragraph(
            doc,
            text,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent_cm=1.0,
            space_after=6,
        )
        content_index += 1

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
